import asyncio
import json
from io import BytesIO

import pytest

from pageindex.core.indexers import DocumentIndexer, IndexerDependencies
from tests.helpers import FakeLLMClient, build_docx_bytes, build_docx_bytes_with_field_table

try:
    from docx import Document  # type: ignore
except Exception:
    Document = None


# The table expander prompt is loaded from this filename — when the prompt
# loader interpolates it, the resulting prompt contains a tag string we can
# match against to know "this is a table-expansion gate request".
_TABLE_GATE_PROMPT_MARKER = "should_expand"


def _run_indexer(path, llm_client=None, **option_overrides):
    options = {
        "if_add_node_id": "yes",
        "if_add_node_summary": "yes",
        "if_add_doc_description": "yes",
        "summary_token_threshold": 1,
    }
    options.update(option_overrides)
    indexer = DocumentIndexer(
        IndexerDependencies(
            libreoffice_command="libreoffice",
            doc_conversion_timeout_seconds=1,
            model="gpt-test",
        )
    )
    return asyncio.run(
        indexer.index(
            file_path=path,
            index_options=options,
            llm_client=llm_client or FakeLLMClient(),
        )
    )


def _llm_client_that_expands_tables(name_column_index: int):
    """Return a FakeLLMClient that says ``should_expand=true`` for table
    gate prompts and falls back to "fake-response" for everything else."""

    def responder(prompt: str, json_response: bool) -> str:
        if json_response and _TABLE_GATE_PROMPT_MARKER in prompt:
            return json.dumps({"should_expand": True, "name_column_index": name_column_index})
        return "fake-response"

    return FakeLLMClient(response_for_prompt=responder)


def _collect_node_ids(structure):
    ids: set[str] = set()
    if isinstance(structure, dict):
        if "node_id" in structure:
            ids.add(structure["node_id"])
        ids.update(_collect_node_ids(structure.get("nodes") or []))
    elif isinstance(structure, list):
        for item in structure:
            ids.update(_collect_node_ids(item))
    return ids


def _collect_leaf_node_ids(structure):
    leaves: set[str] = set()
    if isinstance(structure, dict):
        children = structure.get("nodes") or []
        if not children and "node_id" in structure:
            leaves.add(structure["node_id"])
        leaves.update(_collect_leaf_node_ids(children))
    elif isinstance(structure, list):
        for item in structure:
            leaves.update(_collect_leaf_node_ids(item))
    return leaves


def _require_python_docx():
    if Document is None:
        pytest.skip("python-docx is required for DOCX integration tests")


def test_docx_indexer_builds_tree_and_summary(tmp_path):
    try:
        payload = build_docx_bytes()
    except RuntimeError as exc:
        if "python-docx" in str(exc):
            pytest.skip(str(exc))
        raise

    path = tmp_path / "sample.docx"
    path.write_bytes(payload)

    result = _run_indexer(path)

    assert result["doc_name"] == "sample"
    structure = result["structure"]
    assert any(n["title"] == "Executive Summary" for n in structure)

    # extract.blocks must be present and consistent with char_count
    assert "extract" in result
    blocks = result["extract"]["blocks"]
    assert len(blocks) > 0
    assert sum(len(b["normalized_text"]) for b in blocks) == result["char_count"]
    assert result["location_unit"] == "section"

    structure_node_ids = _collect_node_ids(structure)
    for block in blocks:
        assert block["metadata"]["type"] == "text"
        assert block["metadata"]["pageindex_node_id"] in structure_node_ids


def test_docx_indexer_expands_field_table_when_llm_says_yes(tmp_path):
    """When the LLM gate marks a table as field-definition style, each
    row should become its own outline node and its own block(s), with the
    body text carrying a breadcrumb path so embedding sees the full
    document context."""

    try:
        payload = build_docx_bytes_with_field_table()
    except RuntimeError as exc:
        if "python-docx" in str(exc):
            pytest.skip(str(exc))
        raise

    path = tmp_path / "field-table.docx"
    path.write_bytes(payload)

    # The fixture's table has columns: 分类 / 字段名称 / 类型 / 说明
    # Tell the LLM to use 字段名称 (index 1) as the row name.
    llm_client = _llm_client_that_expands_tables(name_column_index=1)

    result = _run_indexer(
        path,
        llm_client=llm_client,
        if_add_node_id="no",
        if_add_node_summary="no",
        if_add_doc_description="no",
        if_add_node_text="yes",
    )

    # Field rows show up in the structure as their own nodes
    structure_titles = _collect_node_titles(result["structure"])
    assert "核算组织" in structure_titles
    assert "编码" in structure_titles

    # Each field row becomes its own block, and the body block carries the
    # breadcrumb path + full row content (no markdown ## / ###).
    block_texts = [b["normalized_text"] for b in result["extract"]["blocks"]]
    joined = "\n".join(block_texts)

    assert "路径: 租赁物资档案（大卡片） > 核算组织" in joined
    assert "路径: 租赁物资档案（大卡片） > 编码" in joined
    assert "字段名称: 核算组织" in joined
    assert "类型: 组织" in joined
    assert "说明: 当前核算组织" in joined
    assert "字段名称: 编码" in joined
    assert "类型: 文本" in joined
    # The body never resorts to the legacy markdown formatting
    assert "## 基本信息" not in joined
    assert "### 核算组织" not in joined


def test_docx_indexer_falls_back_to_plain_table_when_llm_says_no(tmp_path):
    try:
        payload = build_docx_bytes_with_field_table()
    except RuntimeError as exc:
        if "python-docx" in str(exc):
            pytest.skip(str(exc))
        raise

    path = tmp_path / "field-table-no.docx"
    path.write_bytes(payload)

    # FakeLLMClient default returns "fake-response" → invalid JSON →
    # validator falls back → table goes through plain rendering.
    result = _run_indexer(
        path,
        if_add_node_id="no",
        if_add_node_summary="no",
        if_add_doc_description="no",
        if_add_node_text="yes",
    )

    structure_titles = _collect_node_titles(result["structure"])
    # No per-row nodes — fields are NOT separate sections
    assert "核算组织" not in structure_titles
    assert "编码" not in structure_titles

    # The plain-row rendering still shows up somewhere as a table chunk
    block_texts = "\n".join(b["normalized_text"] for b in result["extract"]["blocks"])
    assert "核算组织" in block_texts  # value is in there, just not as a heading
    assert "编码" in block_texts


def _collect_node_titles(structure):
    titles: list[str] = []
    if isinstance(structure, dict):
        if "title" in structure:
            titles.append(structure["title"])
        titles.extend(_collect_node_titles(structure.get("nodes") or []))
    elif isinstance(structure, list):
        for item in structure:
            titles.extend(_collect_node_titles(item))
    return titles


def test_docx_indexer_emits_blocks_for_no_heading_doc(tmp_path):
    _require_python_docx()
    document = Document()
    document.add_paragraph("First body paragraph.")
    document.add_paragraph("Second body paragraph.")
    document.add_paragraph("Third body paragraph.")
    buffer = BytesIO()
    document.save(buffer)

    path = tmp_path / "no-headings.docx"
    path.write_bytes(buffer.getvalue())

    result = _run_indexer(path, if_add_node_summary="no", if_add_doc_description="no")

    blocks = result["extract"]["blocks"]
    assert len(blocks) == 3  # one block per paragraph

    # All blocks fall into section 1 (the fallback section for headingless docs)
    assert all(b["start_index"] == 1 for b in blocks)
    assert all(b["end_index"] == 1 for b in blocks)
    assert all(b["page_no"] == 1 for b in blocks)

    # Block ordering within the single section is monotonic
    assert [b["block_no"] for b in blocks] == [1, 2, 3]
    assert [b["block_order_in_page"] for b in blocks] == [1, 2, 3]

    # All blocks link to the single fallback node in the structure
    structure_node_ids = _collect_node_ids(result["structure"])
    assert len(structure_node_ids) == 1
    expected_node_id = next(iter(structure_node_ids))
    for block in blocks:
        assert block["metadata"]["pageindex_node_id"] == expected_node_id


def test_docx_indexer_handles_nested_tables(tmp_path):
    _require_python_docx()
    document = Document()
    document.add_heading("Nested Demo", level=1)

    outer = document.add_table(rows=1, cols=1)
    outer_cell = outer.rows[0].cells[0]
    outer_cell.text = "outer-cell-text"
    inner = outer_cell.add_table(rows=1, cols=2)
    inner.rows[0].cells[0].text = "inner-left"
    inner.rows[0].cells[1].text = "inner-right"

    buffer = BytesIO()
    document.save(buffer)
    path = tmp_path / "nested.docx"
    path.write_bytes(buffer.getvalue())

    result = _run_indexer(path, if_add_node_summary="no", if_add_doc_description="no")

    block_texts = "\n".join(b["normalized_text"] for b in result["extract"]["blocks"])

    # The nested table content surfaces somewhere in the textualized blocks.
    assert "inner-left" in block_texts
    assert "inner-right" in block_texts


def test_docx_indexer_preserves_empty_columns_in_plain_tables(tmp_path):
    _require_python_docx()
    document = Document()
    document.add_heading("Table Demo", level=1)
    table = document.add_table(rows=2, cols=3)
    table.rows[0].cells[0].text = "A"
    table.rows[0].cells[1].text = "B"
    table.rows[0].cells[2].text = "C"
    table.rows[1].cells[0].text = "D"
    table.rows[1].cells[1].text = ""
    table.rows[1].cells[2].text = "F"

    buffer = BytesIO()
    document.save(buffer)
    path = tmp_path / "plain-table.docx"
    path.write_bytes(buffer.getvalue())

    result = _run_indexer(path, if_add_node_summary="no", if_add_doc_description="no")

    block_texts = "\n".join(b["normalized_text"] for b in result["extract"]["blocks"])

    # Empty middle cell preserved as a separator slot, so column alignment
    # is recoverable from the textualized table.
    assert "A | B | C" in block_texts
    assert "D |  | F" in block_texts


def test_docx_indexer_block_node_links_resolve_to_covering_node(tmp_path):
    """Each block must link to the deepest tree node whose section range
    covers the block's section_ordinal. That is usually a leaf, but a
    block may legitimately link to an internal parent when the parent
    owns content in a section that none of its children cover (e.g. body
    paragraphs that appear under a heading before its first sub-heading).
    """

    _require_python_docx()
    document = Document()
    # Three sibling leaves at sections 1, 2, 3 — every block lives under a
    # leaf, so we can assert leaf-only linkage in this shape.
    document.add_heading("Section A", level=1)
    document.add_paragraph("Body of A.")
    document.add_heading("Section B", level=1)
    document.add_paragraph("Body of B.")
    document.add_heading("Section C", level=1)
    document.add_paragraph("Body of C.")

    buffer = BytesIO()
    document.save(buffer)
    path = tmp_path / "siblings.docx"
    path.write_bytes(buffer.getvalue())

    result = _run_indexer(path, if_add_node_summary="no", if_add_doc_description="no")

    leaf_node_ids = _collect_leaf_node_ids(result["structure"])
    for block in result["extract"]["blocks"]:
        assert block["metadata"]["pageindex_node_id"] in leaf_node_ids


def test_docx_indexer_block_links_to_parent_when_no_child_covers_section(tmp_path):
    """When a heading has its own body paragraphs before any sub-heading,
    those paragraphs' blocks must link to the parent heading (which is the
    only node whose section range covers them) rather than failing the
    lookup.
    """

    _require_python_docx()
    document = Document()
    document.add_heading("Parent", level=1)
    document.add_paragraph("Parent's own body paragraph.")
    document.add_heading("Child", level=2)
    document.add_paragraph("Child body.")

    buffer = BytesIO()
    document.save(buffer)
    path = tmp_path / "parent-owns-content.docx"
    path.write_bytes(buffer.getvalue())

    result = _run_indexer(path, if_add_node_summary="no", if_add_doc_description="no")

    structure_node_ids = _collect_node_ids(result["structure"])
    for block in result["extract"]["blocks"]:
        # All blocks must link to *some* node — never None — even when the
        # only covering node is an internal parent.
        node_id = block["metadata"]["pageindex_node_id"]
        assert node_id is not None
        assert node_id in structure_node_ids


def test_docx_indexer_block_count_matches_body_items(tmp_path):
    _require_python_docx()
    document = Document()
    document.add_heading("Top", level=1)
    document.add_paragraph("p1")
    document.add_paragraph("p2")
    document.add_heading("Sub", level=2)
    document.add_paragraph("p3")
    document.add_table(rows=1, cols=2).rows[0].cells[0].text = "tablecell"

    buffer = BytesIO()
    document.save(buffer)
    path = tmp_path / "count.docx"
    path.write_bytes(buffer.getvalue())

    result = _run_indexer(path, if_add_node_summary="no", if_add_doc_description="no")

    blocks = result["extract"]["blocks"]
    # 2 headings + 3 paragraphs + 1 table = 6 blocks
    assert len(blocks) == 6
