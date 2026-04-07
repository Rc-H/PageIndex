from __future__ import annotations

from io import BytesIO
from typing import Any

from pageindex.infrastructure.llm import LLMClient

try:
    from docx import Document
except Exception:  # pragma: no cover - optional dependency for test env
    Document = None


class FakeLLMClient(LLMClient):
    def __init__(self, response_for_prompt=None):
        # Optional callable: ``(prompt, json_response) -> str`` so individual
        # tests can return different payloads (e.g. JSON) for specific
        # prompts. When None, falls back to "fake-response".
        self._response_for_prompt = response_for_prompt

    def _resolve_response(self, prompt: str, json_response: bool) -> str:
        if self._response_for_prompt is not None:
            return self._response_for_prompt(prompt, json_response)
        return "fake-response"

    def generate_text(self, model: str, prompt: str, chat_history=None, json_response: bool = False) -> str:
        del model, chat_history
        return self._resolve_response(prompt, json_response)

    def generate_text_with_finish_reason(
        self, model: str, prompt: str, chat_history=None, json_response: bool = False,
    ) -> tuple[str, str]:
        del model, chat_history
        return self._resolve_response(prompt, json_response), "finished"

    async def generate_text_async(
        self, model: str, prompt: str, chat_history=None, json_response: bool = False,
    ) -> str:
        del model, chat_history
        return self._resolve_response(prompt, json_response)


def fake_llm_client_factory() -> LLMClient:
    return FakeLLMClient()


class SpyCallbackClient:
    def __init__(self):
        self.events: list[dict[str, Any]] = []

    async def send(self, callback, payload: dict[str, Any]) -> None:
        del callback
        self.events.append(payload)


class FakeRemoteFileFetcher:
    def __init__(self, file_name: str = "remote.md", content: bytes = b"# Remote\nHello"):
        self.file_name = file_name
        self.content = content

    async def fetch(self, remote_file):
        del remote_file
        from pageindex.messages.models import SubmittedFile

        return SubmittedFile(original_name=self.file_name, content=self.content)


class FakeDocumentIndexer:
    def __init__(self, should_fail: bool = False):
        self.should_fail = should_fail

    async def index(self, file_path, index_options, llm_client):
        del index_options, llm_client
        if self.should_fail:
            raise RuntimeError("index failed")
        return {
            "doc_name": getattr(file_path, "stem", "sample"),
            "structure": [{"title": "Intro", "node_id": "0001"}],
        }


class FakePdfPagePreviewService:
    def __init__(self, previews: list[dict[str, str | int]] | None = None):
        self.previews = previews or []

    def generate(self, file_path):
        del file_path
        return list(self.previews)


def build_docx_bytes() -> bytes:
    if Document is None:
        raise RuntimeError("python-docx is required for DOCX test fixtures")
    document = Document()
    document.add_heading("Executive Summary", level=1)
    document.add_paragraph("Revenue increased year over year.")
    document.add_heading("Details", level=2)
    document.add_paragraph("Operating margin also improved.")
    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def build_docx_bytes_with_field_table() -> bytes:
    if Document is None:
        raise RuntimeError("python-docx is required for DOCX test fixtures")
    document = Document()
    document.add_heading("租赁物资档案（大卡片）", level=1)
    table = document.add_table(rows=3, cols=4)
    table.rows[0].cells[0].text = "分类"
    table.rows[0].cells[1].text = "字段名称"
    table.rows[0].cells[2].text = "类型"
    table.rows[0].cells[3].text = "说明"
    table.rows[1].cells[0].text = "基本信息"
    table.rows[1].cells[1].text = "核算组织"
    table.rows[1].cells[2].text = "组织"
    table.rows[1].cells[3].text = "当前核算组织"
    table.rows[2].cells[0].text = ""
    table.rows[2].cells[1].text = "编码"
    table.rows[2].cells[2].text = "文本"
    table.rows[2].cells[3].text = "自动带出；系统生成"
    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()
